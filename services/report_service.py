"""Report export utilities for DocuFlow-Agent.

This module keeps file generation logic outside the Agent tool layer so the
same Word export capability can be reused by Streamlit UI and Agent tools.
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from utils.logger_handler import logger
from utils.path_tool import get_abs_path


REPORT_PATH_PATTERN = re.compile(r"(?:Word报告已生成|报告已导出|已生成报告)[:：]\s*(?P<path>[^\n]+?\.docx)")


def safe_filename(title: str) -> str:
    """Convert a user-facing title into a filesystem-safe filename stem."""
    filename = re.sub(r"[\\/:*?\"<>|]+", "_", title).strip()
    return filename[:60] or "docuflow_report"


def _iter_markdown_lines(content: str) -> Iterable[str]:
    for line in content.splitlines():
        line = line.strip()
        if line:
            yield line


def export_markdown_to_word(
    content: str,
    title: str = "DocuFlow企业文档分析报告",
    output_dir: str | os.PathLike[str] | None = None,
) -> str:
    """Export markdown-like text to a Word document and return the file path.

    The function intentionally supports a small subset of Markdown because the
    project focuses on reliable enterprise report delivery rather than full
    Markdown rendering.
    """
    if not content.strip():
        raise ValueError("报告正文不能为空")

    report_dir = Path(output_dir) if output_dir else Path(get_abs_path("outputs/reports"))
    report_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filepath = report_dir / f"{safe_filename(title)}_{timestamp}.docx"

    try:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)
        for line in _iter_markdown_lines(content):
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+[\.、]\s*", line):
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)
        doc.save(str(filepath))
        return str(filepath)
    except Exception as exc:  # pragma: no cover - fallback branch
        logger.warning(f"[report_service]生成 Word 失败，降级为 txt：{exc}")
        fallback_path = filepath.with_suffix(".txt")
        fallback_path.write_text(content, encoding="utf-8")
        return str(fallback_path)


def extract_report_path(text: str) -> str | None:
    """Extract a generated report path from Agent output if it exists."""
    match = REPORT_PATH_PATTERN.search(text or "")
    if match:
        return match.group("path").strip()
    return None
