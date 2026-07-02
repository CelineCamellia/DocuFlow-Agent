#文件加载工具的作用是把原始资料文件筛选、读取并转换成 RAG 项目可以处理的 Document 数据，为后续知识库构建和向量检索做准备

import os  # 文件路径判断、目录遍历等操作
import hashlib  # 计算文件 MD5 值
from utils.logger_handler import logger  # 导入项目日志器，用于记录错误信息
from langchain_core.documents import Document  # LangChain 文档对象类型
from langchain_community.document_loaders import PyPDFLoader, TextLoader  # PDF/TXT 文档加载器

# 获取文件的md5的十六进制字符串
def get_file_md5_hex(filepath: str):

    # 判断路径是否存在，不存在则记录错误日志
    if not os.path.exists(filepath):
        logger.error(f"[md5计算]文件{filepath}不存在") #当某个文件路径不存在时，向日志系统记录一条错误信息。
        return

    # 判断路径是否是文件，避免传入文件夹路径
    if not os.path.isfile(filepath):
        logger.error(f"[md5计算]路径{filepath}不是文件") #当路径不是文件时，向日志系统记录一条错误信息。
        return

    md5_obj = hashlib.md5()  # 创建 MD5 计算对象

    chunk_size = 4096  # 4KB分片，避免文件过大爆内存
    try:
        # 以二进制方式读取文件，保证各种文件类型都能参与 MD5 计算
        with open(filepath, "rb") as f:  # 必须二进制读取 #rb:以二进制方式读取文件

            # 分片读取文件内容，并不断更新 MD5 值
            while chunk := f.read(chunk_size): #  := 叫海象运算符，把 f.read(chunk_size) 的结果赋值给chunk，同时把chunk当作while的判断条件
                #读取一块内容给 chunk，如果读到了内容，就继续循环
                md5_obj.update(chunk) #把当前读取到的这一块文件内容加入MD5计算

     #另一种写法：
            """
            chunk = f.read(chunk_size)
            while chunk:

                md5_obj.update(chunk)
                chunk = f.read(chunk_size) #读取下一块内容
            """

            md5_hex = md5_obj.hexdigest()  # 将 MD5 结果转换为十六进制字符串
            return md5_hex
    except Exception as e:
        logger.error(f"计算文件{filepath}md5失败，{str(e)}")  #当计算出错时，向日志系统记录一条错误信息。
        return None

#遍历某个文件夹，只筛选出指定后缀类型的文件，并返回这些文件的完整路径
def listdir_with_allowed_type(path: str, allowed_types: tuple[str]):  # 返回文件夹内的文件列表（允许的文件后缀）
    files = []  # 保存符合要求的文件路径列表

    # 判断传入路径是否是文件夹
    if not os.path.isdir(path):
        logger.error(f"[listdir_with_allowed_type]{path}不是文件夹") #推送错误信息
        return tuple()

    # 遍历文件夹中的文件，筛选指定后缀类型
    for f in os.listdir(path):
        if f.lower().endswith(allowed_types):
            files.append(os.path.join(path, f))  #是符合要求的文件路径，加入文件列表中

    return tuple(files)  # 返回符合要求的文件路径元组，选择元组的原因是，元组不会被轻易改动，所以就是允许获得文件路径，但是不允许修改文件路径

#加载PDF文档
def pdf_loader(filepath: str, passwd=None) -> list[Document]:  #-> list[Document]：返回值类型标注，表示这个函数返回的是list，列表里面每个元素是Document对象。
    # 加载PDF文件，并转换成LangChain的 Document 列表
    return PyPDFLoader(filepath, passwd).load()

#加载文本文档text
def txt_loader(filepath: str) -> list[Document]:
    # 加载 TXT 文件，并转换成 LangChain 的 Document 列表
    return TextLoader(filepath, encoding="utf-8").load()


def docx_loader(filepath: str) -> list[Document]:
    """加载 DOCX 文件，兼容普通段落和表格内容。

    之前版本只读取 doc.paragraphs。很多 Word 文档的主体内容实际在表格、
    编号列表或页眉页脚里，导致“上传成功但知识库没有内容”。这里统一抽取：
    1. 正文段落；
    2. 表格单元格；
    3. 页眉/页脚段落。
    """
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError("请先安装 python-docx：pip install python-docx") from exc

    doc = DocxDocument(filepath)
    blocks: list[str] = []

    # 1. 普通段落，包括 Word 的编号列表文本。
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    # 2. 表格内容。算法题单、简历表格、会议纪要模板经常把正文放在表格里。
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = []
            for cell in row.cells:
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    cells.append(cell_text)
            if cells:
                blocks.append(f"表格{table_index}-行{row_index}: " + " | ".join(cells))

    # 3. 页眉页脚，防止部分模板把关键信息放在这些位置。
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append("页眉: " + text)
        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append("页脚: " + text)

    text = "\n".join(dict.fromkeys(blocks))  # 去掉重复表格单元格文本，保留原顺序。
    if not text.strip():
        logger.warning(f"[docx_loader]{filepath} 未抽取到有效文本，可能是扫描版/图片版 Word")
        return []

    return [Document(page_content=text, metadata={"source": filepath, "file_type": "docx"})]
